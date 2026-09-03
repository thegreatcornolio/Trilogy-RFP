#!/usr/bin/env python3
"""Build the Trilogy CX / Medscheme Master Services Agreement (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x13, 0x20, 0x2E)
GREEN = RGBColor(0x0E, 0x7C, 0x46)
MUTED = RGBColor(0x5B, 0x67, 0x75)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "13202E"
GREEN_HEX = "0E7C46"
ROW_HEX = "F4F7F5"
YELLOW = WD_COLOR_INDEX.YELLOW

OUT = Path(__file__).with_name("Trilogy_CX_Medscheme_MSA.docx")


def set_run_font(run, size=11, bold=False, color=NAVY, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="D0D7DE"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    for child in list(tcPr):
        if child.tag == qn("w:tcBorders"):
            tcPr.remove(child)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, size=10, color=NAVY, align="left", fill=None, highlight=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if highlight:
        run.font.highlight_color = YELLOW
    if fill:
        shade_cell(cell, fill)
    set_cell_borders(cell)
    for inner in cell.paragraphs:
        inner.paragraph_format.line_spacing = 1.08


def add_para(doc, text, *, size=11, bold=False, color=NAVY, space_after=8, space_before=0, align="left"):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_mixed(doc, parts, *, size=11, space_after=8, justify=True):
    """parts: list of str or (str, dict)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    for part in parts:
        if isinstance(part, str):
            run = p.add_run(part)
            set_run_font(run, size=size)
        else:
            text, opts = part
            run = p.add_run(text)
            set_run_font(
                run,
                size=opts.get("size", size),
                bold=opts.get("bold", False),
                color=opts.get("color", NAVY),
            )
            if opts.get("highlight"):
                run.font.highlight_color = YELLOW
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=NAVY, name="Calibri")
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=GREEN)
    else:
        set_run_font(run, size=11.5, bold=True, color=NAVY)
    return p


def page_break(doc):
    doc.add_page_break()


def add_table(doc, headers, rows, col_widths=None, highlight_last=False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9, color=WHITE, fill=NAVY_HEX)
    for r_i, row in enumerate(rows):
        fill = ROW_HEX if r_i % 2 else "FFFFFF"
        is_total = highlight_last and r_i == len(rows) - 1
        for c_i, val in enumerate(row):
            highlight = False
            text = val
            if isinstance(val, tuple):
                text, highlight = val
            set_cell_text(
                table.rows[r_i + 1].cells[c_i],
                text,
                bold=is_total or (c_i == 0 and not str(text).startswith("Week")),
                size=9,
                fill=GREEN_HEX if is_total else fill,
                color=WHITE if is_total else NAVY,
                highlight=highlight,
            )
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def blank_line(doc, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    return p


def sign_block(doc, party_title, names):
    """names: list of (label, value_or_None). None => yellow placeholder."""
    add_para(doc, party_title, size=12, bold=True, space_before=10, space_after=6)
    table = doc.add_table(rows=len(names), cols=2)
    for i, (label, value) in enumerate(names):
        set_cell_text(table.rows[i].cells[0], label, bold=True, size=10, fill=ROW_HEX)
        if value is None:
            set_cell_text(table.rows[i].cells[1], "[to be completed by Medscheme]", size=10, highlight=True)
        else:
            set_cell_text(table.rows[i].cells[1], value, size=10)
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(12.4)
    blank_line(doc, 8)


def footer_header(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.8)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    r = hp.add_run("Trilogy CX (Pty) Ltd  ·  Master Services Agreement  ·  Medscheme")
    set_run_font(r, size=8, bold=True, color=GREEN)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = fp.add_run("Commercial-in-confidence.  Not to be disclosed without the written consent of Trilogy CX (Pty) Ltd.    ")
    set_run_font(r1, size=8, color=MUTED)
    # page number
    r2 = fp.add_run("Page ")
    set_run_font(r2, size=8, color=MUTED)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r3 = fp.add_run()
    r3._r.append(fld1)
    r3._r.append(instr)
    r3._r.append(fld2)
    set_run_font(r3, size=8, color=MUTED)
    r4 = fp.add_run(" of ")
    set_run_font(r4, size=8, color=MUTED)
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    fld4 = OxmlElement("w:fldChar")
    fld4.set(qn("w:fldCharType"), "end")
    r5 = fp.add_run()
    r5._r.append(fld3)
    r5._r.append(instr2)
    r5._r.append(fld4)
    set_run_font(r5, size=8, color=MUTED)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = NAVY
    footer_header(doc)

    add_para(doc, "MASTER SERVICES AGREEMENT", size=22, bold=True, align="center", space_after=4)
    add_para(
        doc,
        "International Business Development and Shared Services Re-Engineering",
        size=12,
        color=GREEN,
        align="center",
        space_after=4,
    )
    add_para(doc, "between Trilogy CX (Pty) Ltd and Medscheme", size=11, align="center", space_after=14, color=MUTED)

    add_mixed(
        doc,
        [
            "This Agreement is dated ",
            ("[insert date of last signature]", {"highlight": True, "bold": True}),
            " (the \"Effective Date\").",
        ],
        space_after=12,
        justify=False,
    )

    heading(doc, "The Parties", 2)
    add_para(doc, "Supplier", size=11, bold=True, space_after=4, color=GREEN)
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Full legal name", "Trilogy CX (Pty) Ltd trading as Trilogy BPO"],
            ["Registration number", "2024/457320/07"],
            ["Physical address", "3rd Floor, Old Mutual Building, Pinelands, Cape Town, 7450"],
            ["Contact person", "Kobus van der Westhuizen"],
            ["Designation", "CEO and Founder"],
            ["E-mail", "kobusvdw@trilogybpo.com"],
        ],
        col_widths=[5.5, 11.2],
    )

    add_para(doc, "Client  (Medscheme to complete the yellow fields)", size=11, bold=True, space_after=4, color=GREEN)
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["Full legal name", ("[insert Medscheme contracting entity]", True)],
            ["Registration number", ("[insert]", True)],
            ["VAT number", ("[insert]", True)],
            ["Physical address", ("[insert]", True)],
            ["Contact person", ("[insert — e.g. Mujeeb Bray]", True)],
            ["Designation", ("[insert — e.g. Chief Operating Officer]", True)],
            ["E-mail", ("[insert]", True)],
            ["Contact number", ("[insert]", True)],
        ],
        col_widths=[5.5, 11.2],
    )

    heading(doc, "1.  Composition of this Agreement", 2)
    add_para(
        doc,
        "1.1  This Agreement comprises the following, which together form one document:",
        space_after=4,
    )
    add_para(doc, "(a)  these cover pages and the signature block;", space_after=2)
    add_para(doc, "(b)  Part A — General Terms and Conditions, which apply to all Services;", space_after=2)
    add_para(
        doc,
        "(c)  Addendum 1 — Sales Enablement Solution (scope, activities, project plan and price in United States Dollars); and",
        space_after=2,
    )
    add_para(
        doc,
        "(d)  Addendum 2 — Shared Services Gate 1 (scope, activities and price in South African Rand).",
        space_after=8,
    )
    add_para(
        doc,
        "1.2  Each Addendum is a statement of work under this Agreement. The same Part A terms apply to both Addenda. Activities, deliverables, term of that workstream and price are set out in the relevant Addendum and differ between them.",
        align="justify",
    )
    add_para(
        doc,
        "1.3  If there is a conflict, the Addendum prevails for the activities and price of that workstream; Part A prevails for legal terms (including payment mechanics, liability, confidentiality, IP, dispute resolution and governing law).",
        align="justify",
    )
    add_para(
        doc,
        "1.4  By signing this Agreement, each party confirms that it has read and understood all parts of this Agreement, and that the person signing is duly authorised.",
        align="justify",
    )

    heading(doc, "2.  Background", 2)
    add_para(
        doc,
        "The Client has healthcare BPO / GBS capability suited to international offshoring. The Supplier will provide (i) international business-development and sales-enablement services into the United Kingdom and United States, working with Precision Point and other consultants engaged by the Supplier, and (ii) a Gate 1 evaluation of the Client’s shared-services function against international BPO best practice. This Agreement records the terms on which those Services are supplied. It replaces the proposal dated 18 / 26 August 2026 once signed, to the extent of any inconsistency.",
        align="justify",
    )

    heading(doc, "3.  Signatures", 2)
    add_para(
        doc,
        "Signed at the places and on the dates set out below. This Agreement may be signed in counterparts (including electronic signature). Each counterpart is an original; together they are one Agreement.",
        align="justify",
        space_after=10,
    )

    sign_block(
        doc,
        "For the Supplier — Trilogy CX (Pty) Ltd",
        [
            ("Name", "Kobus van der Westhuizen"),
            ("Designation", "CEO and Founder"),
            ("Signature", "________________________________"),
            ("Place", "[Cape Town]"),
            ("Date", "[insert]"),
        ],
    )
    sign_block(
        doc,
        "For the Client — Medscheme",
        [
            ("Name", None),
            ("Designation", None),
            ("Signature", "________________________________"),
            ("Place", None),
            ("Date", None),
            ("Who warrants that he/she is duly authorised", ""),
        ],
    )

    # ------------------------------------------------------------------
    page_break(doc)
    heading(doc, "PART A — GENERAL TERMS AND CONDITIONS", 1)
    add_para(
        doc,
        "These terms apply to all Services under this Agreement, including Addendum 1 and Addendum 2. Capitalised terms have the meaning given on the cover pages, in this Part A, or in the relevant Addendum.",
        align="justify",
        color=MUTED,
    )

    heading(doc, "A1.  Definitions", 2)
    defs = [
        ("Addendum", "Addendum 1 or Addendum 2, as the context requires, attached to this Agreement and forming part of it."),
        ("Affiliate", "any person that Controls, is Controlled by, or is under common Control with a party, and \"Control\" has the meaning in the Companies Act, 71 of 2008."),
        ("Charges", "the fees, commission, pass-through costs and other amounts payable by the Client under this Agreement and the applicable Addendum."),
        ("Confidential Information", "all information of a party, in any form, that is by its nature confidential or is designated as such, including client lists, pricing, proposals, personal information, target lists, and the existence and terms of this Agreement, excluding information that is public other than by breach, or independently developed without use of the other party’s information."),
        ("Introduced Client", "any person (a) identified on a written target list prepared under Addendum 1, or (b) to whom the Client, the Supplier or Precision Point makes a sales approach, proposal or introduction in connection with Addendum 1, or (c) with whom the Client contracts for healthcare BPO / GBS or related services where that opportunity arose from the Services."),
        ("Personnel", "directors, officers, employees, contractors and seconded staff of a party or its Affiliates who perform or receive the Services or who have material contact with the other party in connection with this Agreement."),
        ("Personal Information", "has the meaning in the Protection of Personal Information Act, 4 of 2013 (POPIA), and includes personal data under the UK GDPR / EU GDPR where applicable."),
        ("Precision Point", "Precision Point (p3cx.com), engaged by the Supplier as a subcontractor for sales, marketing and market-pricing work under Addendum 1. The Client does not contract with Precision Point."),
        ("Restricted Period", "the Term and a period of 12 (twelve) months after the Term ends, for any reason."),
        ("Services", "the services described in Addendum 1 and Addendum 2, and any further addendum signed by the parties."),
        ("Term", "the period in clause A2, subject to earlier termination in accordance with this Agreement."),
    ]
    for i, (term, meaning) in enumerate(defs, 1):
        add_mixed(
            doc,
            [f"A1.{i}  \"", (term, {"bold": True}), f"\" means {meaning}"],
            space_after=6,
        )

    heading(doc, "A2.  Duration", 2)
    add_para(
        doc,
        "A2.1  This Agreement commences on the Effective Date and continues for 12 (twelve) months, unless terminated earlier under clause A5 (Breach) or clause A2.2. Addendum 2 is a three-month workstream running concurrently from the Effective Date (or such later start date as the parties record in writing). Addendum 1 is a twelve-month workstream from the Effective Date.",
        align="justify",
    )
    add_para(
        doc,
        "A2.2  Either party may terminate this Agreement for no cause on 3 (three) calendar months’ written notice to the other party, such notice not to be given prior to the first anniversary of the Effective Date, provided that this Agreement shall only terminate once all obligations under an Addendum that has a stated minimum term have been fulfilled, as applicable. Termination of one Addendum by written agreement does not of itself terminate the other Addendum or Part A in respect of accrued rights.",
        align="justify",
    )
    add_para(
        doc,
        "A2.3  The expiration or termination of this Agreement shall not affect those provisions which expressly or by necessity survive, including confidentiality, non-solicitation, non-circumvention, data protection, liability, IP, accrued Charges, commission on revenues billed in respect of introductions made during the Term, audit, and dispute resolution.",
        align="justify",
    )

    heading(doc, "A3.  Services", 2)
    add_para(
        doc,
        "A3.1  The Supplier shall provide the Services with reasonable skill and care, in accordance with this Agreement and the applicable Addendum.",
        align="justify",
    )
    add_para(
        doc,
        "A3.2  The Supplier may subcontract Precision Point and other consultants. The Supplier remains responsible to the Client for the Services. The Client shall deal with Precision Point only through the Supplier unless the parties agree otherwise in writing.",
        align="justify",
    )
    add_para(
        doc,
        "A3.3  The Client shall give timely access to information, people, systems and premises reasonably required for the Services, and shall take decisions and grant approvals within the times in the project plan (or otherwise within a reasonable time). Delay by the Client may delay the Services without liability on the Supplier, and Charges remain payable.",
        align="justify",
    )

    heading(doc, "A4.  Charges and payment", 2)
    add_para(
        doc,
        "A4.1  The Client will pay, without set-off or deduction, the Charges specified in the applicable Addendum.",
        align="justify",
    )
    add_para(
        doc,
        "A4.2  Currencies are not combined. Addendum 1 is contracted and billed in United States Dollars (USD). Addendum 2 is contracted and billed in South African Rand (ZAR). Any Rand figure shown against Addendum 1 is indicative only (USD × 18) and is not payable in ZAR unless a ZAR billing arrangement is separately agreed in writing.",
        align="justify",
    )
    add_para(
        doc,
        "A4.3  The Supplier invoices in advance for professional fees. Payment is due within 30 (thirty) days of the invoice date.",
        align="justify",
    )
    add_para(
        doc,
        "A4.4  Pass-through costs (including travel, expenses and marketing under Addendum 1) are billed at cost, in arrears, only where the Client has agreed them in advance (e-mail sufficient). Payment is due within 30 days of invoice.",
        align="justify",
    )
    add_para(
        doc,
        "A4.5  Exchange rates: amounts that are not ZAR, if they must be expressed in another currency for the Client’s internal purposes, are converted at the spot rate at the time of invoice. This does not change the contracting currency of the Addendum.",
        align="justify",
    )
    add_para(
        doc,
        "A4.6  Taxes: all Charges exclude VAT and other taxes, which will be applied at the prevailing rate. Any future tax increases apply but will not be retrospective unless required by law.",
        align="justify",
    )
    add_para(
        doc,
        "A4.7  Late payment: if the Client fails to pay on time, the Supplier may, without prejudice to any other rights: (a) charge interest at Prime Rate (as quoted by the Supplier’s bankers) plus 2%; (b) suspend the Services until full payment is received, the Client remaining liable for Charges during suspension; and (c) terminate this Agreement if payment is not made within ten (10) business days after written notice.",
        align="justify",
    )
    add_para(
        doc,
        "A4.8  Collection costs: the Client shall be liable for all legal costs (on an attorney and own client scale) and collection costs incurred by the Supplier in recovering overdue amounts.",
        align="justify",
    )
    add_para(
        doc,
        "A4.9  Annual increases: after the first anniversary of the Effective Date, professional fees under a continuing or renewed Addendum will increase by 7% unless the parties agree otherwise in writing, on at least 30 days’ notice before the anniversary, taking effect on the anniversary of the Effective Date.",
        align="justify",
    )

    heading(doc, "A5.  Breach", 2)
    add_para(
        doc,
        "If the Client breaches the payment terms and does not remedy the breach within 5 days of its due date, or if either party breaches any other term in this Agreement and fails to remedy such breach within 10 business days of written demand from the other party, or if either party makes a declaration of a moratorium in respect of any indebtedness, or makes or proposes any arrangement or compromise with or any assignment for the benefit of its creditors due to financial difficulties, or ceases to carry on business, or is provisionally or finally wound-up other than a voluntary liquidation for reconstruction or amalgamation, or is placed under business rescue, then the aggrieved party, without prejudice to any other rights it may have, may either demand specific performance or terminate this Agreement and in either case claim damages. If the defaulting party is the Client, then the Supplier may, in addition, suspend the Services until it is satisfied that the breach has been remedied, and the Client shall remain liable to pay the Charges applicable during such suspension.",
        align="justify",
    )

    heading(doc, "A6.  Confidentiality", 2)
    add_para(
        doc,
        "A6.1  Both parties will keep all Confidential Information received from the other party confidential and secure, and will use it only for this Agreement. The receiving party must return or destroy such information upon termination or upon request, save for information it is required to retain for legal reasons or in routine IT backups (which remain subject to this clause until deleted in the ordinary cycle).",
        align="justify",
    )
    add_para(
        doc,
        "A6.2  A receiving party may disclose Confidential Information strictly to its employees, representatives, professional advisors and direct suppliers (including Precision Point, in the case of the Supplier) who are required to know it for this Agreement, provided that such person is subject to confidentiality undertakings no less strict than this clause, or if required to be disclosed by law. If required to be disclosed by law, the receiving party shall notify the disclosing party as soon as possible prior to disclosure, if it is legally possible and/or able.",
        align="justify",
    )
    add_para(
        doc,
        "A6.3  The obligations in this clause A6 survive the Term for five (5) years, and indefinitely in respect of trade secrets and Personal Information.",
        align="justify",
    )

    heading(doc, "A7.  Non-solicitation of personnel", 2)
    add_para(
        doc,
        "A7.1  During the Restricted Period, neither party shall, directly or indirectly, solicit, entice, employ or engage (or attempt to do so) any Personnel of the other party (or of Precision Point, in the case of the Client) who was involved in the Services or with whom that party had material contact in connection with this Agreement, without the other party's prior written consent.",
        align="justify",
    )
    add_para(
        doc,
        "A7.2  This clause does not prevent a party from hiring a person who responds to a bona fide public advertisement not targeted at that person, or whose employment with the other party has ended more than three months before the approach, without prior solicitation in breach of A7.1.",
        align="justify",
    )
    add_para(
        doc,
        "A7.3  If a party breaches clause A7.1, it shall pay the other party, as a genuine pre-estimate of recruitment, training and disruption loss (and not as a penalty), an amount equal to six (6) months' cost-to-company of the person concerned, within 30 days of written demand. This is in addition to any other remedy, including interdict.",
        align="justify",
    )
    add_para(
        doc,
        "A7.4  The parties record that this restraint is reasonable as to time, persons and activity, given access to Confidential Information, client relationships and trained staff. If a court finds any part too wide, that part shall be reduced to the maximum enforceable extent.",
        align="justify",
    )

    heading(doc, "A8.  Non-solicitation of clients and non-circumvention", 2)
    add_para(
        doc,
        "A8.1  During the Restricted Period, the Client shall not circumvent the Supplier by contracting with Precision Point (or any other consultant introduced by the Supplier) for sales, marketing, market-pricing or introductions of the kind in Addendum 1, except through the Supplier or with the Supplier's prior written consent.",
        align="justify",
    )
    add_para(
        doc,
        "A8.2  During the Restricted Period, the Supplier shall not, for its own account or for a competing healthcare BPO / GBS provider, solicit or accept an engagement from an Introduced Client to provide the same class of healthcare administration / BPO services that the Client was proposing to that Introduced Client, except with the Client's prior written consent. This does not prevent the Supplier from providing CX, GCC, digital or other services that are not a substitute for the Client's healthcare BPO offering.",
        align="justify",
    )
    add_para(
        doc,
        "A8.3  Closing a deal with an Introduced Client without the Supplier does not avoid commission. If, during the Restricted Period, the Client (or an Affiliate) bills revenue to an Introduced Client for healthcare BPO / GBS or related services, the commission in Addendum 1 remains payable whether or not the Supplier is still engaged on that account.",
        align="justify",
    )
    add_para(
        doc,
        "A8.4  Nothing in this clause A8 makes the Supplier the Client's exclusive route to all UK/US healthcare demand, except that opportunities on the written target list (and other Introduced Clients) are covered as set out above. The Supplier may serve other clients in other sectors and geographies.",
        align="justify",
    )

    heading(doc, "A9.  Data protection", 2)
    add_para(
        doc,
        "A9.1  Each party shall comply with POPIA and, where it processes Personal Information of persons in the UK or EEA in connection with the Services, with the UK GDPR / EU GDPR as applicable. Gate 1 and sales enablement are not intended to include processing of medical-scheme member records or special personal information. If that processing becomes necessary, the parties will sign a written operator / processor addendum before it starts.",
        align="justify",
    )
    add_para(
        doc,
        "A9.2  Each party is responsible for Personal Information of its own staff. The Client warrants that it has a lawful basis to share with the Supplier any business-contact data of prospects and Introduced Clients. The Supplier shall not sell that data and shall use it only for the Services.",
        align="justify",
    )
    add_para(
        doc,
        "A9.3  A party shall notify the other without undue delay, and in any event within 48 hours of becoming aware, of a security compromise affecting Personal Information received from the other party, and shall reasonably cooperate on any regulator or data-subject response.",
        align="justify",
    )

    heading(doc, "A10.  Insurance", 2)
    add_para(
        doc,
        "A10.1  For the Term and 12 months thereafter, the Supplier shall maintain, with a reputable insurer, public liability and professional indemnity insurance appropriate to the Services, each with a limit of not less than R10,000,000 per claim (or equivalent). The Supplier shall provide evidence of cover on reasonable written request.",
        align="justify",
    )
    add_para(
        doc,
        "A10.2  Insurance does not limit the Supplier's liability except as set out in clause A21.",
        align="justify",
    )

    heading(doc, "A11.  Warranties", 2)
    add_para(
        doc,
        "A11.1  Each party warrants that it is duly incorporated, has power and authority to enter into this Agreement, and that the person signing is authorised.",
        align="justify",
    )
    add_para(
        doc,
        "A11.2  The Supplier warrants that it will perform the Services with reasonable skill and care, using Personnel who are reasonably qualified for the tasks assigned to them, and that it will not knowingly infringe third-party intellectual property in materials it creates.",
        align="justify",
    )
    add_para(
        doc,
        "A11.3  The Client warrants that information it supplies is, to its knowledge, accurate and not misleading in any material respect, and that use of its brands, logos and scheme materials in collateral approved by it will not infringe third-party rights.",
        align="justify",
    )
    add_para(
        doc,
        "A11.4  Except as set out in this Agreement, all other warranties, whether express or implied by law, are excluded to the maximum extent permitted.",
        align="justify",
    )

    heading(doc, "A12.  Indemnity", 2)
    add_para(
        doc,
        "A12.1  The Client indemnifies the Supplier against claims, damages and reasonable legal costs arising from (a) the Client's materials, brands or instructions, (b) the Client's use of deliverables other than as licensed, and (c) any Introduced Client or regulator claim to the extent caused by the Client's operations, pricing, or healthcare administration -- except to the extent caused by the Supplier's negligence or wilful default.",
        align="justify",
    )
    add_para(
        doc,
        "A12.2  The Supplier indemnifies the Client against third-party claims that materials originally created by the Supplier for an Addendum infringe that third party's intellectual property, provided the Client gives prompt notice, does not admit liability, and allows the Supplier to conduct the defence. This indemnity does not apply to Client materials, third-party platforms, or combinations not approved by the Supplier.",
        align="justify",
    )

    heading(doc, "A13.  Relationship of the parties", 2)
    add_para(
        doc,
        "A13.1  The Supplier is an independent contractor. Nothing in this Agreement creates a partnership, joint venture, employment, or agency. Neither party may bind the other or hold itself out as doing so, except that the Supplier and Precision Point may represent that they are mandated to market the Client's international healthcare BPO capability, within approved collateral and messages.",
        align="justify",
    )
    add_para(
        doc,
        "A13.2  The Supplier's Personnel remain its employees or contractors. This Agreement does not transfer any employee to the Client (or the reverse). If any law treats a person as employed by the other party, the first party shall indemnify the other against resulting employment claims, except where that result was caused by the other party's own conduct.",
        align="justify",
    )

    heading(doc, "A14.  Publicity", 2)
    add_para(
        doc,
        "Neither party shall issue a press release or public announcement about this Agreement or name the other in marketing without prior written consent, except that (a) the Supplier may describe the engagement in confidential pitches as a healthcare GBS / BPO mandate, without disclosing commercial terms, and (b) both parties may use approved logos and case language once the Client has signed off collateral under Addendum 1. Consent shall not be unreasonably withheld for a factual announcement of the partnership.",
        align="justify",
    )

    heading(doc, "A15.  Change control", 2)
    add_para(
        doc,
        "A change to scope, deliverables, timetable or Charges is only binding if recorded in a written change note signed by both parties (e-mail with a signed attachment is sufficient). The Supplier is not obliged to perform out-of-scope work. Work done at the Client's written request outside an Addendum may be charged at the Supplier's then-current rates.",
        align="justify",
    )

    heading(doc, "A16.  Records and audit", 2)
    add_para(
        doc,
        "A16.1  The Client shall keep reasonable records of revenues billed to Introduced Clients, sufficient to calculate commission, for the Restricted Period. On not more than one occasion per calendar year (unless a discrepancy of more than 5% was found on the previous review), the Supplier may, on 10 business days' notice, inspect those records during normal hours, itself or through a professional advisor bound by confidentiality. The Client shall cooperate in good faith.",
        align="justify",
    )
    add_para(
        doc,
        "A16.2  If an inspection shows underpayment of more than 5% of commission due for the period reviewed, the Client shall pay the shortfall within 14 days and the reasonable cost of the inspection. Overpayment shall be credited against the next invoice.",
        align="justify",
    )

    heading(doc, "A17.  Anti-bribery and sanctions", 2)
    add_para(
        doc,
        "Each party shall comply with applicable anti-bribery and anti-corruption laws (including the Prevention and Combating of Corrupt Activities Act, 12 of 2004, and, where relevant, the UK Bribery Act 2010) and shall not offer or accept an improper payment or benefit in connection with this Agreement. Each party represents that it is not a sanctioned person and is not owned or controlled by one. A material breach of this clause is a material breach of this Agreement not capable of remedy.",
        align="justify",
    )

    heading(doc, "A18.  Conflicts of interest", 2)
    add_para(
        doc,
        "The Supplier shall disclose in writing any material conflict between this mandate and another client that would reasonably be expected to prejudice the Client's UK/US healthcare BPO campaign. The Supplier may serve other BPO and CX clients. Where a conflict cannot be managed by information barriers and the Client objects on reasonable grounds, the parties shall discuss in good faith whether that workstream should be paused or re-scoped.",
        align="justify",
    )

    heading(doc, "A19.  Key personnel", 2)
    add_para(
        doc,
        "The Supplier's lead for this Agreement is Kobus van der Westhuizen, who may delegate day-to-day delivery. The Supplier shall not permanently remove that lead from the engagement without reasonable notice and a suitable replacement, except for illness, resignation or cause. Precision Point resources are supplied through the Supplier and may change on notice, provided continuity of the plan in Addendum 1 is not unreasonably affected.",
        align="justify",
    )

    heading(doc, "A20.  Return of property", 2)
    add_para(
        doc,
        "On request, and in any event within 15 business days after the Term, each party shall return or securely destroy the other party's property, access credentials and Confidential Information (subject to A6.1). The Client remains licensed to use paid-for deliverables as in clause A22.",
        align="justify",
    )

    heading(doc, "A21.  Liability", 2)
    add_para(
        doc,
        "A21.1  The Supplier shall only be liable for its own actions or omissions and shall not be liable for: any act or omissions of any third party (other than its subcontractors in respect of the Services); unavailability of systems due to disruption of or any failure or fault in apparatus, lines and systems (including a failure in any third party's network); suspension of Services for routine or emergency maintenance or for non-payment; any event or circumstance which is out of the reasonable control of the Supplier (a force majeure event); any failure of non-Supplier equipment; any breach of security which occurs notwithstanding the Supplier's compliance with its security undertakings.",
        align="justify",
    )
    add_para(
        doc,
        "A21.2  The Supplier's liability, howsoever arising, is limited to the value of the professional fees that the Client has paid under the Addendum which is the subject of a claim in the immediate prior 12 months. Neither party is liable for indirect, incidental, or consequential damages, including loss of profit, revenue, goodwill or anticipated savings -- except that this exclusion does not apply to (a) unpaid Charges or commission, (b) a breach of confidentiality, non-solicitation or non-circumvention, or (c) the indemnities in clause A12, which remain subject to A21.3.",
        align="justify",
    )
    add_para(
        doc,
        "A21.3  Nothing in this Agreement excludes liability for fraud, death or personal injury caused by negligence, or any liability that cannot be excluded by South African law.",
        align="justify",
    )

    heading(doc, "A22.  Intellectual property", 2)
    add_para(
        doc,
        "Each party retains ownership of its intellectual property, with no transfer of rights unless expressly agreed in writing. The Client is granted a non-exclusive licence to use deliverables produced under an Addendum for its internal business purposes. The Supplier may reuse generic know-how, methodologies and templates. Marketing and sales collateral produced for the Client may be used by both parties for the purpose of the engagement. Client marks remain the Client's; the Supplier's marks remain the Supplier's.",
        align="justify",
    )

    heading(doc, "A23.  Force majeure", 2)
    add_para(
        doc,
        "The Supplier is not liable for delays or non-performance due to events beyond its control, such as natural disasters, epidemic, utility failure, or strikes. If such events last over three months, the Client may terminate the affected Addendum on 10 business days' written demand.",
        align="justify",
    )

    heading(doc, "A24.  Dispute resolution", 2)
    add_para(
        doc,
        "A24.1  In the event of a dispute arising from or in connection with this Agreement, the parties shall first attempt to resolve the matter through good faith negotiations.",
        align="justify",
    )
    add_para(
        doc,
        "A24.2  If unresolved within 10 business days, the dispute shall be referred to mediation, and if necessary, arbitration in Johannesburg in accordance with the rules of the Arbitration Foundation of Southern Africa (AFSA). This clause shall not preclude either party from seeking urgent interim relief, and the Supplier shall be entitled to approach any court of competent jurisdiction for any unpaid Charges.",
        align="justify",
    )

    heading(doc, "A25.  General", 2)
    generals = [
        "A25.1  This Agreement will be governed under the laws of South Africa, and both parties submit to the exclusive jurisdiction of the South African courts (subject to clause A24).",
        "A25.2  This Agreement is the whole agreement between the Client and the Supplier and excludes any express or implied or tacit term, representation, warranty, promise or the like not recorded herein, whether it induced the contract or not, save for any statutory requirements. Any variation to this Agreement must be in writing and signed between the parties.",
        "A25.3  A failure or delay by a party to exercise any right or remedy hereunder or by law shall not constitute a waiver of that or any other right or remedy.",
        "A25.4  All notices must be delivered by hand or email to the address on the cover pages, or any other address notified to the other party, which takes effect on the 5th day of written notification. Any notice given in writing and actually received by a party's authorised representatives will be deemed properly given, notwithstanding that it has not been given in accordance with this clause.",
        "A25.5  When any number of days is prescribed, same shall be reckoned exclusively of the first and inclusively of the last day unless the last day falls on a Saturday, Sunday or South African public holiday, in which case the last day shall be the next succeeding business day.",
        "A25.6  It is specifically recorded that whilst the parties may correspond via e-mail for operational reasons, no formal legal notice nor any amendment to the Agreement will be of any force and effect if communicated via e-mail, unless specifically for purposes of attaching a signed notice, save that (i) invoices and purchase-order references may be sent by e-mail, and (ii) prior agreement to pass-through costs may be given by e-mail.",
        "A25.7  The provisions of the Agreement are severable and the invalidity of any one or more of such provisions will in no way affect the validity of the remaining provisions. The parties intend clauses A7 and A8 to be enforced to the maximum extent permitted.",
        "A25.8  Neither party may assign or transfer (in whole or in part) any of its rights or obligations under the Agreement, without the other party's prior written consent (such consent not to be unreasonably withheld or delayed), except that the Supplier may assign to an affiliate in the Trilogy group on written notice.",
        "A25.9  The Agreement may be executed in any number of counterparts, including electronic signature contemplated in the Electronic Communications and Transactions Act, 25 of 2002. No counterpart shall be effective until each party has executed at least one counterpart.",
        "A25.10  The provisions of this Agreement shall be binding upon the successors-in-title and assigns of the parties.",
        "A25.11  The rule of interpretation that a contract is to be interpreted against the party responsible for drafting shall not apply.",
        'A25.12  "Calendar month" means a month beginning on the first day of any given month and ending on the last day of the same month. Where notice is given on any day of the month other than the first, the notice period will only commence on the first day of the following month.',
        "A25.13  Each party shall, at its own cost, do all things reasonably required to give effect to this Agreement.",
    ]
    for g in generals:
        add_para(doc, g, align="justify", space_after=6)

    # ------------------------------------------------------------------
    page_break(doc)
    heading(doc, "ADDENDUM 1 — SALES ENABLEMENT SOLUTION", 1)
    add_para(
        doc,
        "This Addendum sets out the activities, project plan and price for the Sales Enablement Solution. Part A applies. Contracting and billing currency: United States Dollars (USD).",
        align="justify",
        color=MUTED,
    )

    heading(doc, "1.1  Purpose", 2)
    add_para(
        doc,
        "The Supplier, working with Precision Point, will position Medscheme International in the UK and US healthcare BPO market — initiation and market positioning, then market engagement and sales — supported by a Cape Town based lead-generation team.",
        align="justify",
    )

    heading(doc, "1.2  Term of this Addendum", 2)
    add_para(
        doc,
        "Twelve (12) months from the Effective Date. Months 1–3 are the initiation and market-positioning phase (the project plan below). Months 4–12 are market engagement and sales, at the fees in clause 1.6.",
        align="justify",
    )

    heading(doc, "1.3  Who does the work", 2)
    add_para(
        doc,
        "Work is led as marked on the project plan: Trilogy-led, Medscheme-led, Precision Point-led, or joint. Initials used on the plan include KvdW (Kobus van der Westhuizen), MB (Mujeeb Bray or such other Client nominee), CS / LM (Supplier project leads), and KE / MR / RF (Precision Point). The Client will nominate its counterparts in writing within five business days of the Effective Date.",
        align="justify",
    )

    heading(doc, "1.4  Activities and project plan — Months 1 to 3", 2)
    add_para(
        doc,
        "The following is the Sales & Marketing Project Schedule from the 18 August 2026 presentation. Weeks are counted from the Effective Date (Week 1 = the first full week). Bars may overlap. Internal MI / Trilogy updates are recurring throughout.",
        align="justify",
        space_after=10,
    )

    heading(doc, "Month 1 — Design & Prepare (Weeks 1–4, with some tasks into Weeks 5–6)", 3)
    add_table(
        doc,
        ["Activity", "Lead", "Timing"],
        [
            ["Introduction of partners", "Trilogy (KvdW)", "Week 2"],
            ["Onboarding & contracting", "Medscheme (CS)", "Weeks 2–4"],
            ["Develop project charter", "Trilogy (CS / LM)", "Weeks 2–4"],
            ["Project planning workshop", "Joint — Trilogy / Medscheme (KvdW / CS / MB / other)", "Weeks 3–4"],
            ["Develop marketing plan", "Joint — Trilogy / Precision Point (KvdW / KE)", "Weeks 3–5"],
            ["Develop sales plan — market engagement", "Joint — Trilogy / Precision Point (KvdW / KE / MR)", "Weeks 3–6"],
            ["Develop market positioning & messages", "Joint — Trilogy / Precision Point (KvdW / MR / MI person)", "Weeks 4–6"],
            ["Develop sales & marketing collateral", "Joint — Trilogy / Medscheme / Precision Point", "Weeks 4–6"],
            ["Identify target companies & executives", "Precision Point (MR / KE / RF)", "Weeks 3–6"],
            ["Internal MI / Trilogy updates", "Joint — Trilogy / Medscheme (KvdW / CS / MB) — recurring", "From Week 2, ongoing"],
        ],
        col_widths=[7.2, 6.0, 3.5],
    )

    heading(doc, "Month 2 — Prepare & Engage (Weeks 5–8)", 3)
    add_table(
        doc,
        ["Activity", "Lead", "Timing"],
        [
            ["Finalise & approve sales & marketing content / collateral", "Joint — Medscheme / Trilogy / Precision Point", "Week 6"],
            ["Finalise & approve marketing plan", "Joint — Medscheme / Trilogy / Precision Point", "Week 6"],
            ["Finalise & approve sales engagement plan", "Joint — Trilogy / Precision Point", "Weeks 6–7"],
            ["Market pricing analysis", "Precision Point", "Weeks 6–7"],
            ["Finalise MI pricing model", "Joint — Medscheme / Trilogy / Precision Point", "Weeks 7–8"],
            ["US partner training workshop", "Joint — Medscheme / Trilogy / Precision Point", "Week 8"],
            ["Proposal / RFP content development", "Joint — Medscheme / Trilogy / Precision Point", "Weeks 7–9"],
            ["Launch social media marketing campaign", "Joint — Trilogy / Precision Point", "Weeks 8–9"],
            ["Internal MI / Trilogy updates", "Joint — Trilogy / Medscheme — recurring", "Weeks 6–9, ongoing"],
        ],
        col_widths=[7.2, 6.0, 3.5],
    )

    heading(doc, "Month 3 — Client meetings (Weeks 9–12)", 3)
    add_table(
        doc,
        ["Activity", "Lead", "Timing"],
        [
            ["Proposal / RFP content development (continued)", "Joint — Medscheme / Trilogy / Precision Point", "Weeks 10–11"],
            ["Scale social media marketing", "Trilogy", "Weeks 10–12"],
            ["Client engagement — per sales plan", "Precision Point", "Weeks 10–12"],
            ["Feedback and re-design of value proposition", "Joint — Precision Point / Trilogy / Medscheme", "Weeks 10–12"],
            ["Internal MI / Trilogy updates", "Joint — Trilogy / Medscheme — recurring", "Weeks 10–12, ongoing"],
            ["Strategy and update-status workshop", "Joint — Precision Point / Trilogy / Medscheme", "Weeks 11–12"],
            ["Develop action plan for Months 4–12", "Joint — Trilogy / Medscheme / Precision Point", "Weeks 10–12"],
        ],
        col_widths=[7.2, 6.0, 3.5],
    )

    heading(doc, "1.5  Months 4–12 — Market engagement & sales", 2)
    add_para(doc, "Unless the parties agree a replacement plan at the Month 3 workshop, Months 4–12 comprise:", align="justify", space_after=4)
    add_para(doc, "(a)  market engagement and sales activity in line with the approved sales plan;", space_after=2)
    add_para(doc, "(b)  continued Cape Town lead-generation team;", space_after=2)
    add_para(doc, "(c)  client meetings, proposals and RFP responses as originated;", space_after=2)
    add_para(doc, "(d)  social and other marketing as approved;", space_after=2)
    add_para(doc, "(e)  recurring internal MI / Trilogy updates; and", space_after=2)
    add_para(
        doc,
        "(f)  validation and, if required, revision of the commission percentage once the retail “Market Pricing” exercise in Months 1–3 is complete.",
        space_after=8,
    )

    heading(doc, "1.6  Professional fees (USD)", 2)
    add_para(
        doc,
        "Year-1 professional fees. Months 4–12 continue at the Month 4–6 run-rate. All amounts exclude VAT.",
        align="justify",
    )
    add_table(
        doc,
        ["Line", "M1", "M2", "M3", "M4", "M5", "M6", "Year 1"],
        [
            ["Initiation & market positioning (Months 1–3)", "$15,000", "$15,000", "$15,000", "—", "—", "—", "$45,000"],
            ["Market engagement & sales (Months 4–12)", "—", "—", "—", "$10,000", "$10,000", "$10,000", "$90,000"],
            ["Travel, expenses & marketing", "Pass-through at cost (not in this fee)", "—", "—", "—", "—", "—", "$0 in this fee"],
            ["Trilogy lead-gen team (Cape Town, Months 1–12)", "$5,000", "$5,000", "$5,000", "$5,000", "$5,000", "$5,000", "$60,000"],
            ["Total professional fees (USD)", "$20,000", "$20,000", "$20,000", "$15,000", "$15,000", "$15,000", "$195,000"],
        ],
        highlight_last=True,
    )
    add_para(
        doc,
        "Year 1 path: Months 1–3 at $20,000 per month; Months 4–12 at $15,000 per month ($10,000 market engagement + $5,000 lead gen). Annual total $45,000 + $90,000 + $60,000 = USD 195,000. Months 7–12 are billed at the same $15,000 monthly run-rate as Months 4–6.",
        align="justify",
        size=10,
        color=MUTED,
    )
    add_para(
        doc,
        "Indicative ZAR (not payable): at USD × 18, Year 1 is approximately R3,510,000. The Client pays this Addendum in USD unless a ZAR billing arrangement is separately agreed.",
        align="justify",
        size=10,
        color=MUTED,
    )

    heading(doc, "1.7  Commission", 2)
    add_para(
        doc,
        "The Client will pay a commission of 15% (fifteen percent) on all revenues billed by the Client directly to clients in respect of work originated or supported under this Addendum, including revenues billed to Introduced Clients as defined in Part A. This percentage will be validated and revised once the Supplier and Precision Point have completed the retail “Market Pricing” exercise in the first three months. Commission is due within 30 days of the Client invoicing the underlying client, and is payable in the same currency as that underlying invoice (or USD if the parties so agree). Survival: commission remains payable on such revenues billed after expiry of this Addendum where the introduction or opportunity arose during the Term, for 12 months after expiry (aligned with the Restricted Period), unless the parties agree otherwise in writing when revising the percentage. The Client shall notify the Supplier in writing within 10 business days of issuing an invoice to an Introduced Client, stating the client name, invoice date and amount (exclusive of VAT).",
        align="justify",
    )

    heading(doc, "1.8  Target list", 2)
    add_para(
        doc,
        "By the end of Week 6 the parties shall maintain a written target list of companies and executives (updated as names are added). A name on that list is an Introduced Client for Part A. Removing a name requires both parties’ written agreement. The list is Confidential Information.",
        align="justify",
    )

    heading(doc, "1.9  Pass-through costs", 2)
    add_para(
        doc,
        "Travel, expenses and marketing are excluded from the professional-fee total and will be billed at cost, with prior agreement (e-mail sufficient).",
        align="justify",
    )

    heading(doc, "1.10  Client dependencies", 2)
    add_para(doc, "The Client shall, without limitation:", space_after=4)
    add_para(doc, "(a)  complete onboarding and contracting in Weeks 2–4;", space_after=2)
    add_para(doc, "(b)  attend the project planning workshop and US partner training workshop;", space_after=2)
    add_para(doc, "(c)  approve marketing plan, sales engagement plan, collateral and the MI pricing model within five business days of submission;", space_after=2)
    add_para(doc, "(d)  make subject-matter experts available for positioning, proposals and client meetings; and", space_after=2)
    add_para(doc, "(e)  not hold the Supplier in delay where a Client-led or joint task is late for reasons on the Client’s side.", space_after=8)

    heading(doc, "1.11  Acknowledgement of this Addendum", 2)
    add_para(
        doc,
        "The signatures on the cover pages apply to this Addendum. If the parties wish to initial this page they may do so below.",
        align="justify",
        space_after=8,
    )
    add_para(doc, "Supplier initial: __________     Client initial: __________     Date: __________", size=10)

    # ------------------------------------------------------------------
    page_break(doc)
    heading(doc, "ADDENDUM 2 — SHARED SERVICES GATE 1", 1)
    add_para(
        doc,
        "This Addendum sets out the activities and price for the Shared Services Gate 1 evaluation. Part A applies. Contracting and billing currency: South African Rand (ZAR).",
        align="justify",
        color=MUTED,
    )

    heading(doc, "2.1  Purpose", 2)
    add_para(
        doc,
        "A three-month evaluation of the Client’s shared-services function against international BPO best practice, so the function can support international commercial models. Centralising recruitment, training, QA and analytics, WFM, MIS and onboarding is intended to give one operating standard, economies of scale, faster scaling and one source of truth — not a different standard at every site.",
        align="justify",
    )

    heading(doc, "2.2  The full model (for context only)", 2)
    add_para(
        doc,
        "The end-state model is a four-phase cycle across every Hub discipline — Discovery, Evaluation, Measurement and Improvement — typically described as six months of advisory and practical augmentation of resources, technology and process. This Addendum does not purchase that full model. It purchases Gate 1 only, being the first three phases. Improvement is a later, separately agreed phase.",
        align="justify",
    )
    add_table(
        doc,
        ["Phase", "In this Addendum?", "What happens"],
        [
            ["1. Discovery", "Yes — Gate 1", "Supplier practitioners embed in the Medscheme environment to observe and document how each Hub discipline operates today."],
            ["2. Evaluation", "Yes — Gate 1", "Each discipline is assessed against international best practice, process maturity and outcomes."],
            ["3. Measurement", "Yes — Gate 1", "Findings are benchmarked against the Supplier’s standards, surfacing gaps and opportunities."],
            ["4. Improvement", "No — out of scope", "A tailored recommendation for the shared-services model best suited to Medscheme. Priced only after Gate 1 is reviewed."],
        ],
        col_widths=[3.8, 3.4, 9.5],
    )

    heading(doc, "2.3  Term of this Addendum", 2)
    add_para(
        doc,
        "Three (3) months from the Effective Date (or such later start date as the parties record in writing). The workstream may run in parallel with Addendum 1.",
        align="justify",
    )

    heading(doc, "2.4  Disciplines in scope", 2)
    add_para(doc, "Gate 1 evaluates the following Hub disciplines:", space_after=4)
    for d in [
        "Recruitment — sourcing and screening talent matched to the Client’s operational needs;",
        "Training — structured induction that gets practitioners productive;",
        "Onboarding — documented, low-risk transition of new starters into live operations;",
        "Quality assurance — continuous quality scoring and root-cause insight;",
        "Workforce management — forecasting and scheduling aligned to demand;",
        "MIS reporting — a single operating view rather than fragmented, delayed reports; and",
        "AI insights & analytics — analysis that surfaces patterns and opportunities.",
    ]:
        add_para(doc, "•  " + d, space_after=2)
    add_para(
        doc,
        "Incubation and continuous improvement sit in the wider Hub model; they are described for context and are not a separate fee line in Gate 1. Commercial contract management (centralised SLA, pricing and commercial-term oversight) is part of the operating picture against which Gate 1 is assessed.",
        align="justify",
        space_before=6,
    )

    heading(doc, "2.5  Activities (Gate 1)", 2)
    add_para(doc, "The Supplier shall, across the disciplines in clause 2.4:", space_after=4)
    add_para(doc, "(a)  Discovery — embed practitioners, observe and document current-state process, roles, tools, volumes and pain points;", space_after=2)
    add_para(doc, "(b)  Evaluation — assess each discipline against international best practice, process maturity and outcomes;", space_after=2)
    add_para(doc, "(c)  Measurement — benchmark findings against the Supplier’s standards and surface gaps and opportunities;", space_after=2)
    add_para(
        doc,
        "(d)  apply the Gate 1 criteria of Innovation, Global Best Practice and BPO-ready. “Areas for augmentation” is the output of Gate 1 (the identified gaps), not a fourth paid phase in this Addendum; and",
        space_after=2,
    )
    add_para(
        doc,
        "(e)  deliver a written roadmap with measurable outcomes and associated cost efficiencies, for the Client’s review.",
        space_after=8,
    )
    add_para(
        doc,
        "Gate 1 is designed to identify: opportunities to reduce overheads through streamlining; gearing shared services to global BPO best practice for international business; and building a capability that can support international commercial BPO models.",
        align="justify",
    )

    heading(doc, "2.6  Out of scope", 2)
    add_para(
        doc,
        "Improvement / augmentation implementation, build of a new hub, software licences, and any work beyond the three-month Gate 1 evaluation, unless agreed in a further addendum. The full four-phase model remains available after Gate 1 on terms to be agreed.",
        align="justify",
    )

    heading(doc, "2.7  Professional fees (ZAR)", 2)
    add_para(doc, "All amounts exclude VAT. This Addendum is billed in Rand. No USD equivalent is proposed.", align="justify")
    add_table(
        doc,
        ["Item", "Term", "Monthly fee", "Total"],
        [
            ["Shared Services Gate 1 — Discovery, Evaluation, Measurement", "3 months", "R195,000", "R585,000"],
            ["Total professional fees (ZAR)", "3 months", "R195,000", "R585,000"],
        ],
        highlight_last=True,
        col_widths=[8.4, 2.6, 2.8, 2.8],
    )

    heading(doc, "2.8  Outcome", 2)
    add_para(
        doc,
        "Improvement recommendations and a roadmap, with measurable outcomes and associated cost efficiencies. Implementation of that roadmap is a separate commercial, agreed after Gate 1 findings are reviewed.",
        align="justify",
    )

    heading(doc, "2.9  Client dependencies", 2)
    add_para(doc, "The Client shall give the Supplier reasonable access to shared-services leaders, process documentation, sample MI, and sites as required for embedding, and shall attend a close-out readout of the roadmap.", align="justify")

    heading(doc, "2.10  Acknowledgement of this Addendum", 2)
    add_para(
        doc,
        "The signatures on the cover pages apply to this Addendum. If the parties wish to initial this page they may do so below.",
        align="justify",
        space_after=8,
    )
    add_para(doc, "Supplier initial: __________     Client initial: __________     Date: __________", size=10)

    add_para(
        doc,
        "End of Master Services Agreement — Trilogy CX (Pty) Ltd and Medscheme.",
        size=9,
        color=MUTED,
        space_before=24,
        align="center",
    )

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    build()
